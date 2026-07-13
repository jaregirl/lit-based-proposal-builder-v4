from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "erb-templates"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def base_document(form_title):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    for name, size in (("Heading 1", 15), ("Heading 2", 12)):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x18, 0x54, 0x4B)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("HOLY NAME UNIVERSITY\n")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(13)
    run2 = title.add_run(form_title)
    run2.bold = True
    run2.font.name = "Arial"
    run2.font.size = Pt(15)

    notice = doc.add_paragraph()
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    notice_run = notice.add_run("DRAFT FOR ADVISER AND ERB REVIEW - NOT APPROVED FOR RECRUITMENT OR DATA COLLECTION")
    notice_run.bold = True
    notice_run.font.name = "Arial"
    notice_run.font.size = Pt(9)
    notice_run.font.color.rgb = RGBColor(0xA3, 0x1D, 0x1D)

    doc.add_paragraph("Template basis: HNU Informed Consent / Parent or LAR Permission / Child Assent Forms v2. Verify the latest official ERB form and submission requirements before use.")
    metadata_table(doc)
    return doc


def metadata_table(doc):
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    labels = [
        ("Study Title", "{{STUDY_TITLE}}"),
        ("Principal Investigator", "{{PRINCIPAL_INVESTIGATOR}}"),
        ("Co-Investigator/s", "{{CO_INVESTIGATOR}}"),
        ("Faculty Adviser", "{{FACULTY_ADVISER}}"),
        ("Study Sponsor", "{{SPONSOR}}"),
    ]
    for row, (label, token) in zip(table.rows, labels):
        row.cells[0].width = Inches(1.65)
        row.cells[1].width = Inches(5.2)
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        row.cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(row.cells[0], "EAF3F0")
        p = row.cells[0].paragraphs[0]
        p.add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(token)


def section(doc, heading, token, lead=""):
    doc.add_heading(heading, level=2)
    p = doc.add_paragraph()
    if lead:
        p.add_run(lead + " ")
    p.add_run(token)


def common_body(doc):
    section(doc, "Purpose and Significance", "{{PURPOSE}}", "You are invited to consider joining this study. Its purpose and significance are described below:")
    section(doc, "Who May Participate", "{{INCLUSION}}", "Inclusion criteria:")
    section(doc, "Who May Not Participate", "{{EXCLUSION}}", "Exclusion criteria:")
    section(doc, "Expected Number of Participants", "{{EXPECTED_PARTICIPANTS}}")
    section(doc, "Research Environment or Setting", "{{SETTING}}")
    section(doc, "What Participation Involves", "{{ACTIVITIES}}")
    section(doc, "Time Commitment", "{{TIME_COMMITMENT}}")
    section(doc, "Risks or Discomforts", "{{RISKS}}")
    section(doc, "Safeguards", "{{SAFEGUARDS}}")
    section(doc, "Possible Benefits", "{{BENEFITS}}")
    section(doc, "Compensation, Reimbursement, or Gift", "{{COMPENSATION}}")
    section(doc, "Privacy and Confidentiality", "{{SAFEGUARDS}}", "Study information will be handled according to the following safeguards:")
    section(doc, "Questions or Concerns", "{{CONTACT}}", "You may contact the researcher through:")
    doc.add_heading("Voluntary Participation and Withdrawal", level=2)
    doc.add_paragraph("Participation is voluntary. Refusal to participate or a decision to withdraw will not result in penalty or loss of benefits to which the participant is otherwise entitled. Ask the researcher to explain anything that is unclear before signing.")


def signature_block(doc, role_label, include_child=False, two_representatives=False):
    doc.add_page_break()
    doc.add_heading("Consent / Permission Record", level=1)
    doc.add_paragraph("I have read or had this information explained to me. I had an opportunity to ask questions. By signing below, I document consent or permission for participation. I understand that this draft must be reviewed and replaced or confirmed by the official ERB-approved form before recruitment or data collection.")
    lines = []
    if two_representatives:
        lines.extend([
            ("First parent / legally authorized representative name", "Signature", "Date"),
            ("Second parent / legally authorized representative name", "Signature", "Date"),
        ])
    else:
        lines.append((role_label, "Signature", "Date"))
    if include_child:
        lines.append(("Child participant name", "Child signature / mark", "Date"))
    lines.append(("Person obtaining consent / permission", "Signature", "Date"))
    table = doc.add_table(rows=len(lines), cols=3)
    table.style = "Table Grid"
    for row, labels in zip(table.rows, lines):
        for cell, label in zip(row.cells, labels):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
            p = cell.paragraphs[0]
            p.add_run("\n\n____________________________\n").bold = True
            p.add_run(label)


def build_adult():
    doc = base_document("INFORMED CONSENT FORM")
    doc.add_paragraph("This form is for an adult participant who is able to provide informed consent.")
    common_body(doc)
    signature_block(doc, "Participant name")
    return doc


def build_parent(two=False):
    title = "TWO-REPRESENTATIVE PARENT / LAR PERMISSION FORM" if two else "PARENT / LEGALLY AUTHORIZED REPRESENTATIVE PERMISSION FORM"
    doc = base_document(title)
    doc.add_paragraph("This form seeks permission for a minor or another participant who requires a parent or legally authorized representative. Child age range: {{CHILD_AGE}}.")
    common_body(doc)
    if two:
        doc.add_paragraph("Two representative signatures are included only because the adviser or ERB explicitly required them. The ERB must determine whether this pathway is appropriate.")
    signature_block(doc, "Parent / legally authorized representative name", two_representatives=two)
    return doc


def build_assent():
    doc = base_document("CHILD ASSENT FORM")
    doc.add_paragraph("This form is written for children in the following age range: {{CHILD_AGE}}. Planned language: {{ASSENT_LANGUAGE}}. Adapt the wording to the child's age, language, and level of understanding, then obtain adviser and ERB review.")
    doc.add_heading("Why are we doing this study?", level=2)
    doc.add_paragraph("{{PURPOSE}}")
    doc.add_heading("What will happen if I join?", level=2)
    doc.add_paragraph("{{ACTIVITIES}}")
    doc.add_heading("How much time will it take?", level=2)
    doc.add_paragraph("{{TIME_COMMITMENT}}")
    doc.add_heading("Could anything feel uncomfortable or difficult?", level=2)
    doc.add_paragraph("{{RISKS}}")
    doc.add_heading("How will the researchers help keep me safe?", level=2)
    doc.add_paragraph("{{SAFEGUARDS}}")
    doc.add_heading("Could this study help anyone?", level=2)
    doc.add_paragraph("{{BENEFITS}}")
    doc.add_heading("Do I have to join?", level=2)
    doc.add_paragraph("No. Joining is your choice. You may say no or stop later. No one should punish you or take away something you should receive because you decide not to join.")
    doc.add_heading("Who can answer my questions?", level=2)
    doc.add_paragraph("{{CONTACT}}")
    signature_block(doc, "Child participant name", include_child=False)
    return doc


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    documents = {
        "adult-consent.docx": build_adult(),
        "parent-permission.docx": build_parent(False),
        "two-representative-permission.docx": build_parent(True),
        "child-assent.docx": build_assent(),
    }
    for name, doc in documents.items():
        doc.save(OUT / name)
        print(OUT / name)


if __name__ == "__main__":
    main()
